from docx import Document
from docx.shared import Pt
from pathlib import Path

path = Path('C:/Users/sunyn/Desktop/project/前端开发工程师-孙一宁-加入AI项目.docx')
doc = Document(path)

# Compress the added AI project to 3 lines if the previous 4-line version exists.
texts_to_remove = {
    '实现多用户 AI 聊天与文件知识库问答，支持模型配置、文件拖拽上传、引用来源预览和 RAG 监控。',
    '接入 Query Rewrite、RAG Router、BM25+向量混合检索、语义切片、多维 metadata 与用户隔离，提升知识库问答召回质量。',
}
merged_line = '项目描述：实现多用户 AI 聊天与文件知识库问答，支持模型配置、文件拖拽上传、流式回复、引用来源预览和 RAG 监控。'
core_line = '核心实现：接入 Query Rewrite、RAG Router、BM25+向量混合检索、语义切片、多维 metadata 与用户隔离。'

paragraphs = list(doc.paragraphs)
for p in paragraphs:
    text = p.text.strip()
    if text in texts_to_remove:
        element = p._element
        element.getparent().remove(element)

# Ensure the compact lines are present after project title/tech stack.
texts = [p.text.strip() for p in doc.paragraphs]
if merged_line not in texts:
    for p in doc.paragraphs:
        if p.text.strip().startswith('技术栈：React、TypeScript'):
            p.insert_paragraph_before(core_line)
            p.insert_paragraph_before(merged_line)
            break

# Reorder compact lines if needed by rebuilding around title would be risky; keep acceptable compact insertion.
for p in doc.paragraphs:
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.95
    for run in p.runs:
        if run.font.size is None or run.font.size.pt > 9:
            run.font.size = Pt(9)

for p in doc.paragraphs:
    if p.text.strip() in {'AI Chat RAG 知识库问答系统', '项目经历', '实习经历', '专业技能', '教育背景', '自我评价'}:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

for section in doc.sections:
    section.top_margin = Pt(28)
    section.bottom_margin = Pt(28)
    section.left_margin = Pt(45)
    section.right_margin = Pt(45)

doc.save(path)
print(path)
