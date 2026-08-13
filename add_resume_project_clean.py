from docx import Document
from docx.shared import Pt
from pathlib import Path

src = Path('C:/Users/sunyn/Desktop/project/前端开发工程师-孙一宁.docx')
out = Path('C:/Users/sunyn/Desktop/project/前端开发工程师-孙一宁-加入AI项目.docx')

doc = Document(src)

project_lines = [
    'AI Chat RAG 知识库问答系统',
    '技术栈：React、TypeScript、Vite、Ant Design、FastAPI、MySQL、Redis、Chroma、Docker。',
    '项目描述：实现多用户 AI 聊天与文件知识库问答，支持模型配置、文件拖拽上传、流式回复、引用来源预览和 RAG 监控。',
    '核心实现：接入 Query Rewrite、RAG Router、BM25+向量混合检索、语义切片、多维 metadata 与用户隔离。',
]

anchor = None
for p in doc.paragraphs:
    if p.text.strip() == '自我评价':
        anchor = p
        break

if anchor is None:
    anchor = doc.add_paragraph('自我评价')

inserted = []
for text in project_lines:
    p = anchor.insert_paragraph_before(text)
    inserted.append(p)

# Compact spacing globally to keep the resume within one page.
for p in doc.paragraphs:
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 0.92
    for run in p.runs:
        if run.font.size is None or run.font.size.pt > 9:
            run.font.size = Pt(9)

# Keep section headings readable.
for p in doc.paragraphs:
    if p.text.strip() in {'教育背景', '专业技能', '实习经历', '项目经历', '自我评价'}:
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

for p in inserted:
    for run in p.runs:
        run.font.size = Pt(8.5)

for run in inserted[0].runs:
    run.bold = True
    run.font.size = Pt(9.5)

# Reduce margins a little, still readable on A4/Word.
for section in doc.sections:
    section.top_margin = Pt(24)
    section.bottom_margin = Pt(24)
    section.left_margin = Pt(42)
    section.right_margin = Pt(42)

doc.save(out)
print(out)
