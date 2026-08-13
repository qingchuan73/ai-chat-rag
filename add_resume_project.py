from docx import Document
from docx.shared import Pt
from pathlib import Path

src = Path('C:/Users/sunyn/Desktop/project/前端开发工程师-孙一宁.docx')
out = Path('C:/Users/sunyn/Desktop/project/前端开发工程师-孙一宁-加入AI项目.docx')

doc = Document(src)

project_lines = [
    'AI Chat RAG 知识库问答系统',
    '技术栈：React、TypeScript、Vite、Ant Design、FastAPI、MySQL、Redis、Chroma、Docker。',
    '实现多用户 AI 聊天与文件知识库问答，支持模型配置、文件拖拽上传、流式回复、引用来源预览和 RAG 监控。',
    '接入 Query Rewrite、RAG Router、BM25+向量混合检索、语义切片、多维 metadata 与用户隔离，提升知识库问答召回质量。',
]

# Insert before self evaluation, keeping existing TodoMVC project.
insert_at = None
for i, p in enumerate(doc.paragraphs):
    if p.text.strip() == '自我评价':
        insert_at = i
        break

if insert_at is None:
    insert_at = len(doc.paragraphs)

anchor = doc.paragraphs[insert_at]
for text in reversed(project_lines):
    p = anchor.insert_paragraph_before(text)
    # Match compact resume style.
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1
    for run in p.runs:
        run.font.size = Pt(9)

# Make project title a little stronger without increasing height.
for p in doc.paragraphs:
    if p.text.strip() == 'AI Chat RAG 知识库问答系统':
        for run in p.runs:
            run.bold = True
            run.font.size = Pt(10)

# Slightly compact all paragraph spacing to help keep one page.
for p in doc.paragraphs:
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    if p.paragraph_format.line_spacing is None:
        p.paragraph_format.line_spacing = 1

# Keep body readable but compact.
for p in doc.paragraphs:
    for run in p.runs:
        if run.font.size is None:
            run.font.size = Pt(9.5)

doc.save(out)
print(out)
